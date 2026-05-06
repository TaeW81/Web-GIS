import io
import requests
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from pyproj import Transformer
from shapely.ops import transform
from config import VWORLD_KEY, VWORLD_DOMAIN, OWNER_CATEGORIES, CADASTRAL_LAYER

def _owner_group(owner_str):
    """소유자 문자열을 (대분류, 소분류)로 매핑"""
    for group, members in OWNER_CATEGORIES.items():
        for m in members:
            if m in str(owner_str):
                return group, owner_str
    return "기타", owner_str

def _set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    
    # HWPX 호환: 한글 폰트 강제 매핑
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rFonts.set(qn('w:ascii'), '맑은 고딕')
    rFonts.set(qn('w:hAnsi'), '맑은 고딕')
    
    if bold: run.bold = True
    if color: run.font.color.rgb = color

def _fmt_area(v):
    try: return f"{float(v):,.1f}"
    except: return str(v)

def _fmt_int(v):
    try: return f"{int(v):,}"
    except: return str(v)

def _fmt_pct(v):
    try: return f"{float(v):.1f}"
    except: return str(v)

class LandReportGenerator:
    """토지이용 현황 보고서 생성기 (고급형 - 복구 버전)"""
    HEADER_COLOR = "4472C4"

    def __init__(self, analysis_data, boundary_polygon=None, gps_points=None,
                 owner_map_bytes=None, jimok_map_bytes=None, location_map_bytes=None):
        self.data = analysis_data or []
        self.boundary_polygon = boundary_polygon
        self.gps_points = gps_points or []
        self.boundary_area = self._calc_boundary_area()
        self.owner_map_bytes = owner_map_bytes
        self.jimok_map_bytes = jimok_map_bytes
        self.location_map_bytes = location_map_bytes
        self.doc = Document()
        
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        self._setup_styles()

    def _calc_boundary_area(self):
        if not self.boundary_polygon: return 0.0
        try:
            proj = Transformer.from_crs("EPSG:4326", "EPSG:5179", always_xy=True).transform
            poly_m = transform(proj, self.boundary_polygon)
            return poly_m.area
        except: return 0.0

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = '맑은 고딕'
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:ascii'), '맑은 고딕')
        rFonts.set(qn('w:hAnsi'), '맑은 고딕')
        style.font.size = Pt(10)
        
        # HWPX 호환: 기본 여백 명시
        section = self.doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _add_run(self, p, text, size=10, bold=False, color=None):
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:ascii'), '맑은 고딕')
        rFonts.set(qn('w:hAnsi'), '맑은 고딕')
        run.font.size = Pt(size)
        run.bold = bold
        if color: run.font.color.rgb = color
        return run

    def generate(self):
        # 1. 위치도 섹션
        self.add_location_section()
        # 2. 소유자별 현황 섹션
        self.add_owner_section()
        # 3. 지목별 현황 섹션
        self.add_jimok_section()
        
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ─────────────────────────────────────
    # 1. 위치도 섹션
    # ─────────────────────────────────────
    def add_location_section(self):
        p_title = self.doc.add_paragraph()
        self._add_run(p_title, "1. 위치도", size=14, bold=True)
        
        p_fig = self.doc.add_paragraph()
        self._add_run(p_fig, "[그림 1-1 위치도]", size=9, bold=True)
        
        img_bytes = self.location_map_bytes or self._render_location_map()
        if img_bytes:
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_bytes, width=Cm(15))
            self._set_table_borders(table)

    def _render_location_map(self):
        if not self.boundary_polygon: return None
        try:
            import matplotlib
            matplotlib.use('Agg')
            from matplotlib.font_manager import FontProperties
            font_prop = FontProperties(fname="C:/Windows/Fonts/malgun.ttf")
            
            centroid = self.boundary_polygon.centroid
            to_utmk = Transformer.from_crs("EPSG:4326", "EPSG:5179", always_xy=True).transform
            cx_m, cy_m = to_utmk(centroid.x, centroid.y)
            
            margin_m = 12500
            to_epsg3857 = Transformer.from_crs("EPSG:5179", "EPSG:3857", always_xy=True).transform
            b_minx, b_miny = to_epsg3857(cx_m - margin_m, cy_m - margin_m)
            b_maxx, b_maxy = to_epsg3857(cx_m + margin_m, cy_m + margin_m)
            
            params_bg = {
                "key": VWORLD_KEY, "service": "image", "request": "getimage",
                "center": f"{centroid.x},{centroid.y}", "level": "11", "size": "900,900"
            }
            params_wms = {
                "service": "WMS", "request": "GetMap", "layers": "lt_c_lhblpn", "styles": "lt_c_lhblpn",
                "crs": "EPSG:3857", "bbox": f"{b_minx},{b_miny},{b_maxx},{b_maxy}",
                "width": "900", "height": "900", "format": "image/png", "transparent": "true",
                "key": VWORLD_KEY, "domain": VWORLD_DOMAIN, "version": "1.3.0"
            }
            
            res_bg = requests.get("http://api.vworld.kr/req/image", params=params_bg, timeout=30)
            bg_img = Image.open(io.BytesIO(res_bg.content)).crop((0, 0, 900, 860))
            
            res_wms = requests.get("http://api.vworld.kr/req/wms", params=params_wms, timeout=30)
            wms_img = None
            if res_wms.status_code == 200 and len(res_wms.content) > 1000:
                wms_img = Image.open(io.BytesIO(res_wms.content)).crop((0, 0, 900, 860))
            
            fig, ax = plt.subplots(figsize=(8, 8))
            img_margin = 11250
            ax.imshow(bg_img, extent=[-img_margin, img_margin, -img_margin + (img_margin*80/900), img_margin])
            if wms_img:
                ax.imshow(wms_img, extent=[-margin_m, margin_m, -margin_m + (margin_m*80/900), margin_m], alpha=0.6)
            
            for radius, label in [(5000, "5km"), (10000, "10km")]:
                circle = plt.Circle((0, 0), radius, color='white', fill=False, linestyle='--', lw=1.5, alpha=0.9)
                ax.add_artist(circle)
                ax.text(radius + 200, 200, label, color='white', fontsize=11, fontweight='bold', fontproperties=font_prop)

            poly_m = transform(to_utmk, self.boundary_polygon)
            for poly in ([poly_m] if poly_m.geom_type == 'Polygon' else poly_m.geoms):
                px, py = poly.exterior.xy
                ax.fill([x - cx_m for x in px], [y - cy_m for y in py], color='yellow', ec='red', lw=2, alpha=1.0, zorder=50)

            ax.set_xlim(-margin_m, margin_m)
            ax.set_ylim(-margin_m, margin_m)
            ax.axis('off')
            buf = io.BytesIO()
            fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', pad_inches=0)
            plt.close(fig)
            buf.seek(0)
            return buf
        except: return None

    # ─────────────────────────────────────
    # 2. 소유자별 현황
    # ─────────────────────────────────────
    def add_owner_section(self):
        self.doc.add_page_break()
        agg = self._aggregate_by_owner()
        ba = self.boundary_area
        total_area = sum(v["area"] for v in agg.values())
        total_count = sum(v["count"] for v in agg.values())
        error_area = ba - total_area if ba > 0 else 0

        p_title = self.doc.add_paragraph()
        self._add_run(p_title, "2-1. 소유자별 현황", size=14, bold=True)
        
        p_tbl_title = self.doc.add_paragraph()
        self._add_run(p_tbl_title, "[표 2-1 소유자별 현황]", size=9, bold=True)
        
        table = self.doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_header_row(table)
        self._add_merged_data_row(table, "구역계 면적", ba, 100.0, total_count, "계 = 면적오차")
        self._add_merged_data_row(table, "면적오차", error_area, (error_area/ba*100 if ba else 0), "-")
        self._add_merged_data_row(table, "계", total_area, (total_area/ba*100 if ba else 0), total_count)
        
        sorted_items = sorted(agg.items(), key=lambda x: x[1]["area"], reverse=True)
        for (grp, sub), v in sorted_items:
            pct = (v["area"] / ba * 100) if ba else 0
            self._add_split_data_row(table, grp, sub, v["area"], pct, v["count"])
        
        self._set_table_borders(table)
        
        p_fig = self.doc.add_paragraph()
        self._add_run(p_fig, "\n[그림 2-1 소유자별 현황도]", size=9, bold=True)
        if self.owner_map_bytes:
            box = self.doc.add_table(rows=1, cols=1)
            box.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = box.cell(0, 0).paragraphs[0].add_run()
            run.add_picture(self.owner_map_bytes, width=Cm(12))
            self._set_table_borders(box)

    # ─────────────────────────────────────
    # 3. 지목별 현황
    # ─────────────────────────────────────
    def add_jimok_section(self):
        self.doc.add_page_break()
        agg = self._aggregate_by_jimok()
        ba = self.boundary_area
        total_area = sum(v["area"] for v in agg.values())
        total_count = sum(v["count"] for v in agg.values())
        error_area = ba - total_area if ba > 0 else 0

        p_title = self.doc.add_paragraph()
        self._add_run(p_title, "2-2. 지목별 현황", size=14, bold=True)
        
        p_tbl_title = self.doc.add_paragraph()
        self._add_run(p_tbl_title, "[표 2-2 지목별 현황]", size=9, bold=True)
        
        table = self.doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 헤더
        headers = ["지목", "면적(㎡)", "비율(%)", "필지수", "비고"]
        for i, h in enumerate(headers):
            _set_cell_text(table.cell(0, i), h, bold=True)
            _set_cell_bg(table.cell(0, i), self.HEADER_COLOR)
        
        # 요약 행
        r1 = table.add_row()
        _set_cell_text(r1.cells[0], "구역계 면적", bold=True); _set_cell_text(r1.cells[1], _fmt_area(ba)); _set_cell_text(r1.cells[2], "100.0"); _set_cell_text(r1.cells[3], _fmt_int(total_count))
        r2 = table.add_row()
        _set_cell_text(r2.cells[0], "면적오차", bold=True); _set_cell_text(r2.cells[1], _fmt_area(error_area)); _set_cell_text(r2.cells[2], _fmt_pct(error_area/ba*100 if ba else 0))
        r3 = table.add_row()
        _set_cell_text(r3.cells[0], "계", bold=True); _set_cell_text(r3.cells[1], _fmt_area(total_area)); _set_cell_text(r3.cells[2], _fmt_pct(total_area/ba*100 if ba else 0)); _set_cell_text(r3.cells[3], _fmt_int(total_count))
        
        sorted_jimok = sorted(agg.items(), key=lambda x: x[1]["area"], reverse=True)
        for jimok, v in sorted_jimok:
            row = table.add_row()
            _set_cell_text(row.cells[0], jimok)
            _set_cell_text(row.cells[1], _fmt_area(v["area"]))
            _set_cell_text(row.cells[2], _fmt_pct(v["area"]/ba*100 if ba else 0))
            _set_cell_text(row.cells[3], _fmt_int(v["count"]))
            
        self._set_table_borders(table)
        
        p_fig = self.doc.add_paragraph()
        self._add_run(p_fig, "\n[그림 2-2 지목별 현황도]", size=9, bold=True)
        if self.jimok_map_bytes:
            box = self.doc.add_table(rows=1, cols=1)
            box.alignment = WD_TABLE_ALIGNMENT.CENTER
            run = box.cell(0, 0).paragraphs[0].add_run()
            run.add_picture(self.jimok_map_bytes, width=Cm(12))
            self._set_table_borders(box)

    def _aggregate_by_owner(self):
        agg = {}
        for p in self.data:
            attr = p.get("analysis_attr", {})
            owner_str = attr.get("소유자", "기타")
            grp, sub = _owner_group(owner_str)
            key = (grp, sub)
            if key not in agg: agg[key] = {"area": 0.0, "count": 0}
            agg[key]["area"] += float(attr.get("편입면적(㎡)", 0.0))
            agg[key]["count"] += 1
        return agg

    def _aggregate_by_jimok(self):
        agg = {}
        for p in self.data:
            attr = p.get("analysis_attr", {})
            jimok = attr.get("지목", "기타")
            if jimok not in agg: agg[jimok] = {"area": 0.0, "count": 0}
            agg[jimok]["area"] += float(attr.get("편입면적(㎡)", 0.0))
            agg[jimok]["count"] += 1
        return agg

    def _add_header_row(self, table):
        headers = ["구분", "세부구분", "면적(㎡)", "비율(%)", "필지수", "비고"]
        row = table.rows[0]
        for i, h in enumerate(headers):
            _set_cell_text(row.cells[i], h, bold=True)
            _set_cell_bg(row.cells[i], self.HEADER_COLOR)

    def _add_merged_data_row(self, table, label, area, pct, count, note=""):
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[2], _fmt_area(area))
        _set_cell_text(row.cells[3], _fmt_pct(pct))
        _set_cell_text(row.cells[4], _fmt_int(count))
        _set_cell_text(row.cells[5], note)

    def _add_split_data_row(self, table, grp, sub, area, pct, count):
        row = table.add_row()
        _set_cell_text(row.cells[0], grp)
        _set_cell_text(row.cells[1], sub)
        _set_cell_text(row.cells[2], _fmt_area(area))
        _set_cell_text(row.cells[3], _fmt_pct(pct))
        _set_cell_text(row.cells[4], _fmt_int(count))

    def _set_table_borders(self, table):
        table.autofit = False
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # HWPX 호환: 표 너비 100% (5000 in pct) 강제 지정
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), "5000")
        tblW.set(qn('w:type'), "pct")
        tblPr.append(tblW)

        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>')
        tblPr.append(borders)
