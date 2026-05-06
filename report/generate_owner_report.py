"""
토지이용현황 - 소유자별 현황 Word 보고서 생성기
첨부된 이미지 양식을 정확히 재현합니다.

구조:
  - 헤더: ② 토지이용현황 | 위치정보 (파란 하단선)
  - 소제목: 2-1. 소유자별
  - 요약 불릿 텍스트
  - [표 2-1 소유자별 현황] 테이블 (6열: 구분(대), 구분(소), 면적, 구성비, 필지수, 비고)
  - 주석
  - [그림 2-1 소유자별 현황도] + 이미지
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


# ===== 샘플 데이터 =====
SAMPLE_DATA = {
    "location": "강원특별자치도 태백시 통동 일원",
    "report_title": "현황분석 보고서",
    "total_area": 2513149.2,
    "total_parcels": 81,
    "area_diff": -13438.9,
    "owners": [
        {"type": "개인", "category": "사유지", "area": 224616.1, "ratio": 8.9, "parcels": 67},
        {"type": "법인", "category": "사유지", "area": 7399.0, "ratio": 0.3, "parcels": 1},
        {"type": "국유지", "category": "국공유지", "area": 2294573.0, "ratio": 91.3, "parcels": 13},
    ],
}


# ===== 유틸리티 함수 =====
def _set_shading(cell, color_hex):
    """셀 배경색"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def _set_borders(cell, **edges):
    """셀 테두리. edges 예: top={"sz":"4","color":"000000"}"""
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, cfg in edges.items():
        borders.append(parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{cfg.get("val","single")}" '
            f'w:sz="{cfg.get("sz","4")}" w:space="0" '
            f'w:color="{cfg.get("color","000000")}"/>'
        ))
    cell._tc.get_or_add_tcPr().append(borders)


def _cell_text(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
               color=None, font_name="맑은 고딕"):
    """셀에 서식 있는 텍스트 입력"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_run(paragraph, text, size=10, bold=False, color=None, font_name="맑은 고딕"):
    """문단에 Run 추가"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _table_full_width(table):
    """테이블 폭 100%"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))


def _no_border(cell):
    """셀 테두리 없음"""
    _set_borders(cell,
                 top={"val": "none", "sz": "0", "color": "FFFFFF"},
                 bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                 left={"val": "none", "sz": "0", "color": "FFFFFF"},
                 right={"val": "none", "sz": "0", "color": "FFFFFF"})


def _thin_border(cell):
    """기본 얇은 테두리"""
    b = {"val": "single", "sz": "4", "color": "000000"}
    _set_borders(cell, top=b, bottom=b, left=b, right=b)


def _merge_cells(table, row1, col1, row2, col2):
    """셀 병합"""
    table.cell(row1, col1).merge(table.cell(row2, col2))


# ===== 메인 생성 함수 =====
def create_owner_report(data=None, output_path=None):
    if data is None:
        data = SAMPLE_DATA

    doc = Document()

    # 페이지 설정 (A4)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 기본 스타일
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    owners = data["owners"]
    total_area = data["total_area"]
    total_parcels = data["total_parcels"]
    area_diff = data["area_diff"]
    sum_area = total_area + abs(area_diff)  # 계 = 구역계 + |면적오차|

    # ============================================================
    # 1) 헤더: ② 토지이용현황          위치 현황분석 보고서
    # ============================================================
    ht = doc.add_table(rows=1, cols=2)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_full_width(ht)

    lc = ht.cell(0, 0)
    lc.text = ""
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(lp, "② ", size=14, bold=True, color=(0, 0, 255))
    _add_run(lp, "토지이용현황", size=14, bold=True)

    rc = ht.cell(0, 1)
    rc.text = ""
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(rp, f"{data['location']} {data['report_title']}", size=9, color=(100, 100, 100))

    for c in ht.rows[0].cells:
        _set_borders(c,
                     top={"val": "none", "sz": "0", "color": "FFFFFF"},
                     left={"val": "none", "sz": "0", "color": "FFFFFF"},
                     right={"val": "none", "sz": "0", "color": "FFFFFF"},
                     bottom={"val": "single", "sz": "12", "color": "0000FF"})

    # ============================================================
    # 2) 소제목
    # ============================================================
    doc.add_paragraph()
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(6)
    p_sub.paragraph_format.space_after = Pt(10)
    _add_run(p_sub, "2-1. 소유자별", size=12, bold=True)

    # ============================================================
    # 3) 요약 불릿
    # ============================================================
    # 요약 불릿 구성
    parts = [f"{o['type']} {o['area']:,.1f}㎡({o['ratio']}%)" for o in owners]
    # 면적 기준 내림차순 정렬
    sorted_owners = sorted(owners, key=lambda x: x['area'], reverse=True)
    sorted_parts = [f"{o['type']} {o['area']:,.1f}㎡({o['ratio']}%)" for o in sorted_owners]

    bullet1 = "소유자별 현황은 " + ", ".join(sorted_parts) + " 순으로 나타남"

    parcel_parts = [f"{o['type']} {o['parcels']}필지" for o in owners]
    bullet2 = "필지수로는 " + ", ".join(parcel_parts) + "로 구성되어 있음"

    bullets = [bullet1, bullet2]

    for txt in bullets:
        if not txt:
            continue
        pb = doc.add_paragraph()
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.left_indent = Cm(0.5)
        _add_run(pb, f"◦ {txt}", size=9)

    doc.add_paragraph()

    # ============================================================
    # 4) 표 제목
    # ============================================================
    pt_ = doc.add_paragraph()
    pt_.paragraph_format.space_after = Pt(4)
    _add_run(pt_, "[표 2-1 소유자별 현황]", size=9, bold=True)

    # ============================================================
    # 5) 소유자별 현황 테이블 (6열 × 7행)
    #    구분(대) | 구분(소) | 면적(㎡) | 구성비(%) | 필지수(필지) | 비고
    # ============================================================
    NUM_ROWS = 7  # 헤더 + 구역계 + 면적오차 + 계 + 개인 + 법인 + 국유지
    NUM_COLS = 6
    tbl = doc.add_table(rows=NUM_ROWS, cols=NUM_COLS)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_full_width(tbl)

    # --- 헤더 (row 0) ---
    # 구분은 col 0~1 병합
    _merge_cells(tbl, 0, 0, 0, 1)
    header_texts = {0: "구분", 2: "면적(㎡)", 3: "구성비(%)", 4: "필지수(필지)", 5: "비고"}
    for ci, txt in header_texts.items():
        c = tbl.cell(0, ci)
        _set_shading(c, "FFFF00")
        _cell_text(c, txt, size=9, bold=True)
        _thin_border(c)

    # --- row 1: 구역계 면적 (col0~1 병합) ---
    _merge_cells(tbl, 1, 0, 1, 1)
    _cell_text(tbl.cell(1, 0), "구역계 면적", size=9)
    _cell_text(tbl.cell(1, 2), f"{total_area:,.1f}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(1, 3), "100.0", size=9)
    _cell_text(tbl.cell(1, 4), str(total_parcels), size=9)
    _cell_text(tbl.cell(1, 5), "제 = 면적오차", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(1, ci))

    # --- row 2: 면적오차 (col0~1 병합) ---
    _merge_cells(tbl, 2, 0, 2, 1)
    _cell_text(tbl.cell(2, 0), "면적오차", size=9)
    _cell_text(tbl.cell(2, 2), f"{area_diff:,.1f}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(2, 3), f"{area_diff / total_area * 100:.1f}", size=9)
    _cell_text(tbl.cell(2, 4), "-", size=9)
    _cell_text(tbl.cell(2, 5), "", size=9)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(2, ci))

    # --- row 3: 계 (col0~1 병합) ---
    _merge_cells(tbl, 3, 0, 3, 1)
    _cell_text(tbl.cell(3, 0), "계", size=9, bold=True)
    _cell_text(tbl.cell(3, 2), f"{sum_area:,.1f}", size=9, bold=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(3, 3), f"{sum_area / total_area * 100:.1f}", size=9, bold=True)
    _cell_text(tbl.cell(3, 4), str(total_parcels), size=9, bold=True)
    _cell_text(tbl.cell(3, 5), "", size=9)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(3, ci))

    # --- row 4: 사유지 - 개인 ---
    # col0 = "사유지" (row4~5 세로병합)
    _merge_cells(tbl, 4, 0, 5, 0)
    _cell_text(tbl.cell(4, 0), "사유지", size=9)
    _cell_text(tbl.cell(4, 1), "개인", size=9)
    _cell_text(tbl.cell(4, 2), f"{owners[0]['area']:,.1f}", size=9,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(4, 3), str(owners[0]['ratio']), size=9)
    _cell_text(tbl.cell(4, 4), str(owners[0]['parcels']), size=9)
    _cell_text(tbl.cell(4, 5), "", size=9)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(4, ci))

    # --- row 5: 사유지 - 법인 ---
    _cell_text(tbl.cell(5, 1), "법인", size=9)
    _cell_text(tbl.cell(5, 2), f"{owners[1]['area']:,.1f}", size=9,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(5, 3), str(owners[1]['ratio']), size=9)
    _cell_text(tbl.cell(5, 4), str(owners[1]['parcels']), size=9)
    _cell_text(tbl.cell(5, 5), "", size=9)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(5, ci))

    # --- row 6: 국공유지 - 국유지 ---
    _cell_text(tbl.cell(6, 0), "국공유지", size=9)
    _cell_text(tbl.cell(6, 1), "국유지", size=9)
    _cell_text(tbl.cell(6, 2), f"{owners[2]['area']:,.1f}", size=9,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(tbl.cell(6, 3), str(owners[2]['ratio']), size=9)
    _cell_text(tbl.cell(6, 4), str(owners[2]['parcels']), size=9)
    _cell_text(tbl.cell(6, 5), "", size=9)
    for ci in range(NUM_COLS):
        _thin_border(tbl.cell(6, ci))

    # ============================================================
    # 6) 주석
    # ============================================================
    notes = [
        '※ 계 면적 산정방식=완전편입 필지(공부면적), 부분편입 필지(구적면적)',
        '※ 토지대장상 정보가 누락되었거나 값이 0인 경우 "기타(값 없음 등)"으로 분류',
    ]
    for nt in notes:
        pn = doc.add_paragraph()
        pn.paragraph_format.space_before = Pt(1)
        pn.paragraph_format.space_after = Pt(1)
        _add_run(pn, nt, size=7, color=(0, 0, 255))

    doc.add_paragraph()

    # ============================================================
    # 7) 그림 제목
    # ============================================================
    pf = doc.add_paragraph()
    pf.paragraph_format.space_after = Pt(6)
    _add_run(pf, "[그림 2-1 소유자별 현황도]", size=9, bold=True)

    # ============================================================
    # 8) 현황도 이미지 (있으면 삽입, 없으면 플레이스홀더)
    # ============================================================
    map_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "map_owner.png")
    if os.path.exists(map_path):
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(map_path, width=Cm(14))
    else:
        pt2 = doc.add_table(rows=1, cols=1)
        pt2.alignment = WD_TABLE_ALIGNMENT.CENTER
        pc = pt2.cell(0, 0)
        pc.height = Cm(10)
        _cell_text(pc, "[소유자별 현황도 이미지 삽입 위치]", size=11, color=(150, 150, 150))
        _set_borders(pc,
                     top={"val": "dashed", "sz": "4", "color": "AAAAAA"},
                     bottom={"val": "dashed", "sz": "4", "color": "AAAAAA"},
                     left={"val": "dashed", "sz": "4", "color": "AAAAAA"},
                     right={"val": "dashed", "sz": "4", "color": "AAAAAA"})

    # 저장
    if output_path is None:
        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..",
                                   "2-1_소유자별_현황.docx")
    output_path = os.path.abspath(output_path)
    doc.save(output_path)
    print(f"✅ 보고서 생성 완료: {output_path}")
    return output_path


if __name__ == "__main__":
    create_owner_report()
